"""DOCX export service for final paper-style documents."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCX_OUTPUT_DIR = PROJECT_ROOT / "outputs" / "docx"


def _ensure_python_docx():
    try:
        import docx  # noqa: F401
        return
    except ModuleNotFoundError:
        pass

    home = Path.home()
    runtime_roots = sorted(home.glob(".cache/codex-runtimes/*/dependencies/python"), reverse=True)
    candidates: list[Path] = []
    for runtime_root in runtime_roots:
        candidates.append(runtime_root / "Lib" / "site-packages")
        candidates.append(runtime_root / "Lib")
        candidates.append(runtime_root)

    for candidate in candidates:
        if not candidate.exists():
            continue
        candidate_text = str(candidate)
        if candidate_text not in sys.path:
            sys.path.insert(0, candidate_text)
        try:
            import docx  # noqa: F401
            return
        except ModuleNotFoundError:
            continue

    raise ModuleNotFoundError("python-docx could not be imported from the local venv or bundled runtime.")


_ensure_python_docx()

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Inches, Pt  # type: ignore


def ensure_output_dir() -> Path:
    DOCX_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return DOCX_OUTPUT_DIR


def normalize_text(text: str) -> str:
    return " ".join(str(text or "").replace("\u00a0", " ").split())


def slugify_topic(topic: str) -> str:
    compact = "_".join((topic or "").strip().lower().split())
    safe = "".join(char if char.isalnum() or char == "_" else "_" for char in compact)
    while "__" in safe:
        safe = safe.replace("__", "_")
    return safe.strip("_") or "report"


def infer_topic_from_report_path(report_path: Path) -> str:
    stem = report_path.stem.replace("_visualized", "")
    stem = re.sub(r"_rev\d+$", "", stem)
    match = re.match(r"(.+?)_\d{8}_\d{6}(?:_\d+)?$", stem)
    if match:
        return match.group(1).replace("_", " ")
    return stem.replace("_", " ")


def find_latest_visualized_report(topic: str | None = None) -> Path | None:
    reports_dir = PROJECT_ROOT / "outputs" / "reports"
    candidates = list(reports_dir.glob("*_visualized.md"))
    if topic:
        prefix = f"{slugify_topic(topic)}_"
        candidates = [path for path in candidates if path.name.startswith(prefix)]
    if not candidates:
        return None
    return max(candidates, key=lambda path: path.stat().st_mtime)


def find_latest_report(topic: str | None = None) -> Path | None:
    reports_dir = PROJECT_ROOT / "outputs" / "reports"
    candidates = [path for path in reports_dir.glob("*.md") if not path.name.endswith("_visualized.md")]
    if topic:
        prefix = f"{slugify_topic(topic)}_"
        candidates = [path for path in candidates if path.name.startswith(prefix)]
    if not candidates:
        return None
    return max(candidates, key=lambda path: path.stat().st_mtime)


def parse_markdown_blocks(markdown_text: str) -> list[dict]:
    blocks: list[dict] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = "\n".join(paragraph_lines).strip()
            if text:
                blocks.append({"type": "paragraph", "text": text})
            paragraph_lines = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        if stripped == "---":
            flush_paragraph()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": heading_match.group(2).strip(),
                }
            )
            continue

        image_match = re.match(r"^!\[[^\]]*\]\((.+)\)$", stripped)
        if image_match:
            flush_paragraph()
            blocks.append({"type": "image", "path": image_match.group(1).strip()})
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def _set_run_font(run, size_pt: int, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 10, line_spacing: float = 1.5) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def apply_paper_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Malgun Gothic"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_title(document: Document, title_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title_text)
    _set_run_font(run, 20, bold=True)
    _set_paragraph_spacing(paragraph, after=16, line_spacing=1.2)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {min(max(level, 1), 3)}"
    run = paragraph.add_run(text)
    _set_run_font(run, 16 if level == 1 else 13, bold=True)
    _set_paragraph_spacing(paragraph, before=8, after=8, line_spacing=1.3)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for index, line in enumerate(text.split("\n")):
        run = paragraph.add_run(line.strip())
        _set_run_font(run, 11, bold=False)
        if index < len(text.split("\n")) - 1:
            run.add_break()
    _set_paragraph_spacing(paragraph, after=10, line_spacing=1.5)


def add_image(document: Document, image_path: Path) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    _set_paragraph_spacing(paragraph, before=6, after=10, line_spacing=1.0)


def export_markdown_to_docx(report_path: Path, topic: str | None = None) -> Path:
    markdown_text = report_path.read_text(encoding="utf-8")
    blocks = parse_markdown_blocks(markdown_text)
    active_topic = topic or infer_topic_from_report_path(report_path)

    document = Document()
    apply_paper_styles(document)

    skip_first_paragraph = False
    if len(blocks) >= 2 and blocks[0].get("type") == "heading" and blocks[1].get("type") == "paragraph":
        heading_text = normalize_text(blocks[0].get("text", "")).lower()
        if heading_text in {"제목", "title"}:
            add_title(document, normalize_text(blocks[1].get("text", "")))
            skip_first_paragraph = True

    for index, block in enumerate(blocks):
        block_type = block.get("type")
        if block_type == "heading":
            heading_text = normalize_text(block.get("text", ""))
            if index == 0 and heading_text.lower() not in {"제목", "title"}:
                add_title(document, heading_text)
            elif heading_text.lower() not in {"제목", "title"}:
                add_heading(document, heading_text, int(block.get("level", 1)))
        elif block_type == "paragraph":
            if skip_first_paragraph:
                skip_first_paragraph = False
                continue
            add_body_paragraph(document, block.get("text", ""))
        elif block_type == "image":
            image_path = Path(block.get("path", ""))
            if not image_path.is_absolute():
                image_path = (report_path.parent / image_path).resolve()
            if image_path.exists():
                add_image(document, image_path)

    ensure_output_dir()
    output_path = DOCX_OUTPUT_DIR / f"{slugify_topic(active_topic)}_{report_path.stem}.docx"
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Export a markdown report to DOCX.")
    parser.add_argument("report_path", help="Absolute path to the markdown report file")
    parser.add_argument("--topic", default="", help="Optional topic override")
    args = parser.parse_args()

    report_path = Path(args.report_path)
    topic = args.topic.strip() or None
    output_path = export_markdown_to_docx(report_path, topic=topic)
    print(f"DOCX_PATH={output_path}")


if __name__ == "__main__":
    main()
